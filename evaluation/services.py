from __future__ import annotations

import hashlib
from dataclasses import dataclass
from io import BytesIO

from django.core.exceptions import ValidationError
from django.core.files.uploadedfile import SimpleUploadedFile
from django.db import transaction
from docx import Document as DocxDocument

from accounts.models import OrganizationMembership, User
from candidates.ai_extraction import (
    CANDIDATE_PROFILE_SCHEMA_VERSION,
    CandidateProfileExtraction,
    SkillEvidence,
    confirm_candidate_profile,
)
from candidates.documents import DOCX_CONTENT_TYPE, upload_candidate_cv
from candidates.models import Candidate, CandidateProfile, CandidateSource
from candidates.services import create_candidate_with_source
from evaluation.dataset import (
    CandidateSpec,
    EvaluationDataset,
    VacancySpec,
    canonical_dataset_json,
)
from matching.models import MatchRun
from matching.scoring import generate_shortlist
from organizations.models import Organization
from vacancies.models import Vacancy, VacancyRequirements
from vacancies.services import (
    change_vacancy_status,
    confirm_requirements_draft,
    create_vacancy_with_requirements,
    update_requirements_draft,
)


class EvaluationDatasetMismatchError(ValidationError):
    """The installed deterministic results differ from the frozen expectation."""


@dataclass(frozen=True)
class InstalledEvaluationDataset:
    organization: Organization
    candidates: dict[str, Candidate]
    vacancies: dict[str, Vacancy]
    match_runs: dict[str, MatchRun]
    dataset_sha256: str


def _build_synthetic_cv(candidate: CandidateSpec) -> bytes:
    document = DocxDocument()
    document.core_properties.title = f"Synthetic evaluation CV {candidate.code}"
    document.core_properties.author = "AI Candidate Matcher evaluation fixture"
    document.add_heading(candidate.full_name, level=0)
    document.add_paragraph(
        "SYNTHETIC EVALUATION FIXTURE - This document describes no real person."
    )
    document.add_paragraph(f"Candidate code: {candidate.code}")
    document.add_heading("Profile", level=1)
    document.add_paragraph(candidate.summary)
    document.add_heading("Location", level=1)
    document.add_paragraph(f"Synthetic location: {candidate.location}.")
    document.add_heading("Skills", level=1)
    document.add_paragraph(
        "Recorded synthetic skills: " + ", ".join(candidate.skills) + "."
    )
    document.add_paragraph(
        "Generated only for controlled AI Candidate Matcher evaluation."
    )
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _profile_extraction(candidate: CandidateSpec) -> CandidateProfileExtraction:
    skill_evidence = "Recorded synthetic skills: " + ", ".join(candidate.skills) + "."
    location_evidence = f"Synthetic location: {candidate.location}."
    return CandidateProfileExtraction(
        relevant_experience_summary=candidate.summary,
        relevant_experience_summary_evidence=candidate.summary,
        skills=[
            SkillEvidence(name=skill, evidence=skill_evidence)
            for skill in candidate.skills
        ],
        location=candidate.location,
        location_evidence=location_evidence,
        ambiguities=[
            "This is a controlled synthetic fixture; unrecorded facts remain unknown."
        ],
    )


def _create_candidate(
    *,
    organization: Organization,
    user: User,
    spec: CandidateSpec,
) -> tuple[Candidate, str]:
    candidate = create_candidate_with_source(
        organization=organization,
        user=user,
        candidate_values={
            "full_name": spec.full_name,
            "email": "",
            "phone": "",
            "location": spec.location,
            "retention_until": None,
        },
        source_values={
            "source_type": CandidateSource.SourceType.OTHER,
            "source_name": "Synthetic EVAL-001 fixture",
            "source_reference": f"EVAL-001-{spec.code}",
            "lawful_basis": CandidateSource.LawfulBasis.NOT_RECORDED,
            "consent_status": CandidateSource.ConsentStatus.NOT_REQUIRED,
            "contact_permission": CandidateSource.ContactPermission.RESTRICTED,
            "permission_notes": (
                "Entirely synthetic evaluation data; outreach is intentionally "
                "restricted."
            ),
            "retention_until": None,
        },
    )
    raw = _build_synthetic_cv(spec)
    document = upload_candidate_cv(
        candidate=candidate,
        user=user,
        uploaded_file=SimpleUploadedFile(
            f"{spec.code.casefold()}-synthetic-evaluation-cv.docx",
            raw,
            content_type=DOCX_CONTENT_TYPE,
        ),
    )
    extraction = _profile_extraction(spec)
    profile = CandidateProfile.objects.create(
        candidate=candidate,
        source_document=document,
        version=1,
        schema_version=CANDIDATE_PROFILE_SCHEMA_VERSION,
        source_document_sha256=document.sha256,
        source_text_sha256=hashlib.sha256(
            document.extracted_text.encode("utf-8")
        ).hexdigest(),
        created_by=user,
        **extraction.as_profile_values(),
    )
    confirm_candidate_profile(profile=profile, user=user)
    return candidate, document.file.name


def _create_vacancy(
    *,
    organization: Organization,
    user: User,
    spec: VacancySpec,
) -> Vacancy:
    vacancy = create_vacancy_with_requirements(
        organization=organization,
        user=user,
        vacancy_values={
            "title": spec.title,
            "description": spec.description,
        },
    )
    draft = vacancy.requirement_versions.get(version=1)
    update_requirements_draft(
        requirements=draft,
        user=user,
        values={
            "summary": spec.summary,
            "must_have_skills": list(spec.must_have_skills),
            "nice_to_have_skills": list(spec.nice_to_have_skills),
            "minimum_years_experience": None,
            "location_requirement": "",
            "work_mode": VacancyRequirements.WorkMode.UNKNOWN,
            "language_requirements": [],
            "education_requirements": [],
            "certification_requirements": [],
            "employment_type": VacancyRequirements.EmploymentType.UNKNOWN,
            "hard_constraints": [],
            "ambiguities": [
                "This vacancy is an entirely synthetic evaluation fixture."
            ],
        },
    )
    confirm_requirements_draft(requirements=draft, user=user)
    change_vacancy_status(
        vacancy=vacancy,
        user=user,
        new_status=Vacancy.Status.OPEN,
    )
    return vacancy


def _assert_expected_ranking(
    *,
    spec: VacancySpec,
    run: MatchRun,
    candidate_codes_by_id: dict[int, str],
) -> None:
    expected = [(item.candidate_code, item.score) for item in spec.expected_top]
    actual = [
        (candidate_codes_by_id[entry.candidate_id], entry.score)
        for entry in run.entries.order_by("rank")[: len(expected)]
    ]
    if actual != expected:
        raise EvaluationDatasetMismatchError(
            f"{spec.code} produced {actual!r}; expected {expected!r}. "
            "No evaluation organization was saved."
        )


@transaction.atomic
def _install_dataset(
    *,
    dataset: EvaluationDataset,
    user: User,
    organization_slug: str,
    stored_files: list[str],
) -> InstalledEvaluationDataset:
    if not user.is_active:
        raise ValidationError("The evaluation dataset requires an active user.")
    if Organization.objects.filter(slug=organization_slug).exists():
        raise ValidationError(
            "An organization already uses this slug. Choose a new evaluation slug; "
            "existing data is never overwritten."
        )

    organization = Organization(
        name=f"Synthetic Evaluation — {dataset.dataset_id}",
        slug=organization_slug,
    )
    organization.full_clean()
    organization.save()
    OrganizationMembership.objects.create(
        user=user,
        organization=organization,
        role=OrganizationMembership.Role.RECRUITER,
    )

    candidates: dict[str, Candidate] = {}
    for spec in dataset.candidates:
        candidate, stored_name = _create_candidate(
            organization=organization,
            user=user,
            spec=spec,
        )
        stored_files.append(stored_name)
        candidates[spec.code] = candidate

    vacancies: dict[str, Vacancy] = {}
    match_runs: dict[str, MatchRun] = {}
    candidate_codes_by_id = {
        candidate.pk: code for code, candidate in candidates.items()
    }
    for spec in dataset.vacancies:
        vacancy = _create_vacancy(
            organization=organization,
            user=user,
            spec=spec,
        )
        requirements = vacancy.current_requirements
        run = generate_shortlist(requirements=requirements, user=user)
        _assert_expected_ranking(
            spec=spec,
            run=run,
            candidate_codes_by_id=candidate_codes_by_id,
        )
        vacancies[spec.code] = vacancy
        match_runs[spec.code] = run

    dataset_sha256 = hashlib.sha256(
        canonical_dataset_json(dataset).encode("utf-8")
    ).hexdigest()
    return InstalledEvaluationDataset(
        organization=organization,
        candidates=candidates,
        vacancies=vacancies,
        match_runs=match_runs,
        dataset_sha256=dataset_sha256,
    )


def install_evaluation_dataset(
    *,
    dataset: EvaluationDataset,
    user: User,
    organization_slug: str,
) -> InstalledEvaluationDataset:
    """Install one isolated synthetic workspace and verify its frozen rankings."""
    stored_files: list[str] = []
    try:
        return _install_dataset(
            dataset=dataset,
            user=user,
            organization_slug=organization_slug,
            stored_files=stored_files,
        )
    except Exception:
        from django.core.files.storage import default_storage

        for stored_name in stored_files:
            default_storage.delete(stored_name)
        raise

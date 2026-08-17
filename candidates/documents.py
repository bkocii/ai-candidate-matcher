import hashlib
import stat
import unicodedata
import zipfile
from dataclasses import dataclass
from datetime import date
from io import BytesIO
from pathlib import PurePosixPath
from xml.etree import ElementTree

from django.core.files.base import ContentFile
from django.db import transaction
from django.utils import timezone
from docx import Document as DocxDocument
from pypdf import PdfReader
from pypdf.generic import ArrayObject, DictionaryObject, IndirectObject

from accounts.models import User
from candidates.models import Candidate, CandidateDocument
from organizations.permissions import require_organization_object_access

DEFAULT_MAX_DOCUMENT_BYTES = 10 * 1024 * 1024
DEFAULT_MAX_PDF_PAGES = 100
DEFAULT_MAX_DOCX_ENTRIES = 2_000
DEFAULT_MAX_DOCX_EXPANDED_BYTES = 50 * 1024 * 1024
DEFAULT_MAX_DOCX_ENTRY_BYTES = 20 * 1024 * 1024
DEFAULT_MAX_EXTRACTED_CHARACTERS = 1_000_000
MAX_DOCX_COMPRESSION_RATIO = 1_000
MAX_PDF_OBJECTS = 25_000
MAX_PDF_OBJECT_DEPTH = 100

PDF_CONTENT_TYPE = "application/pdf"
DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
GENERIC_UPLOAD_CONTENT_TYPES = {"", "application/octet-stream"}


class CandidateDocumentUploadError(ValueError):
    """A safe, recruiter-visible CV upload failure."""

    def __init__(self, code: str, public_message: str):
        self.code = code
        self.public_message = public_message
        super().__init__(public_message)


class CandidateDocumentDuplicateError(CandidateDocumentUploadError):
    """The organization already stores the exact document bytes."""


class CandidateDocumentDeliveryError(ValueError):
    """A safe, recruiter-visible private document delivery failure."""

    def __init__(self, code: str, public_message: str):
        self.code = code
        self.public_message = public_message
        super().__init__(public_message)


@dataclass(frozen=True)
class ExtractedCV:
    text: str
    content_type: str
    extension: str


@dataclass(frozen=True)
class ValidatedCVUpload:
    raw: bytes
    original_filename: str
    extracted: ExtractedCV
    sha256: str


@dataclass(frozen=True)
class PrivateCandidateDocument:
    content: bytes
    filename: str
    content_type: str


def _safe_original_filename(filename: str) -> str:
    cleaned = filename.replace("\\", "/").rsplit("/", maxsplit=1)[-1].strip()
    if not cleaned or any(
        unicodedata.category(character) in {"Cc", "Cf"} for character in cleaned
    ):
        raise CandidateDocumentUploadError(
            "invalid_filename",
            "The uploaded document must have a valid filename.",
        )
    if len(cleaned) > 255:
        raise CandidateDocumentUploadError(
            "filename_too_long",
            "The uploaded filename is longer than 255 characters.",
        )
    return cleaned


def _validate_pdf_object_graph(reader: PdfReader) -> None:
    """Reject executable or embedded PDF features without blocking ordinary links."""
    stack: list[tuple[object, int]] = [(reader.trailer, 0)]
    seen_indirect: set[tuple[int, int]] = set()
    seen_containers: set[int] = set()
    visited = 0

    while stack:
        value, depth = stack.pop()
        if depth > MAX_PDF_OBJECT_DEPTH:
            raise CandidateDocumentUploadError(
                "pdf_structure_too_complex",
                "The PDF structure exceeds the supported safety limit.",
            )

        if isinstance(value, IndirectObject):
            object_key = (value.idnum, value.generation)
            if object_key in seen_indirect:
                continue
            seen_indirect.add(object_key)
            value = value.get_object()

        if not isinstance(value, (DictionaryObject, ArrayObject)):
            continue

        container_id = id(value)
        if container_id in seen_containers:
            continue
        seen_containers.add(container_id)
        visited += 1
        if visited > MAX_PDF_OBJECTS:
            raise CandidateDocumentUploadError(
                "pdf_structure_too_complex",
                "The PDF structure exceeds the supported safety limit.",
            )

        if isinstance(value, DictionaryObject):
            names = {str(key) for key in value}
            if names.intersection({"/JS", "/JavaScript", "/EmbeddedFiles", "/EF"}):
                raise CandidateDocumentUploadError(
                    "unsafe_pdf_content",
                    "PDF scripts, embedded files, and launch actions are not "
                    "supported.",
                )
            if str(value.get("/S", "")) in {
                "/JavaScript",
                "/Launch",
                "/GoToR",
                "/GoToE",
                "/ImportData",
                "/SubmitForm",
            }:
                raise CandidateDocumentUploadError(
                    "unsafe_pdf_content",
                    "PDF scripts, embedded files, and launch actions are not "
                    "supported.",
                )
            if str(value.get("/Type", "")) == "/EmbeddedFile" or str(
                value.get("/Subtype", "")
            ) in {
                "/FileAttachment",
                "/RichMedia",
                "/Movie",
                "/Sound",
            }:
                raise CandidateDocumentUploadError(
                    "unsafe_pdf_content",
                    "PDF scripts, embedded files, and launch actions are not "
                    "supported.",
                )
            stack.extend((item, depth + 1) for item in value.values())
        else:
            stack.extend((item, depth + 1) for item in value)


def _read_bounded(uploaded_file, max_bytes: int) -> bytes:
    chunks: list[bytes] = []
    total = 0
    while total <= max_bytes:
        chunk = uploaded_file.read(max_bytes + 1 - total)
        if not chunk:
            break
        chunks.append(chunk)
        total += len(chunk)
    raw = b"".join(chunks)
    if not raw:
        raise CandidateDocumentUploadError(
            "empty_file",
            "The uploaded document is empty.",
        )
    if len(raw) > max_bytes:
        raise CandidateDocumentUploadError(
            "file_too_large",
            f"The document exceeds the {max_bytes // (1024 * 1024)} MB limit.",
        )
    return raw


def _normalize_extracted_text(text: str, max_characters: int) -> str:
    normalized = "\n".join(
        line.rstrip() for line in text.replace("\x00", "").splitlines()
    ).strip()
    if not normalized:
        raise CandidateDocumentUploadError(
            "no_extractable_text",
            "No readable text was found. Scanned-image CVs are not supported yet.",
        )
    if len(normalized) > max_characters:
        raise CandidateDocumentUploadError(
            "extracted_text_too_large",
            "The extracted document text exceeds the supported limit.",
        )
    return normalized


def _extract_pdf_text(
    raw: bytes,
    *,
    max_pages: int,
    max_characters: int,
) -> str:
    if not raw.startswith(b"%PDF-"):
        raise CandidateDocumentUploadError(
            "invalid_pdf_signature",
            "The file content is not a valid PDF document.",
        )

    try:
        reader = PdfReader(BytesIO(raw), strict=True)
        if reader.is_encrypted:
            raise CandidateDocumentUploadError(
                "encrypted_pdf",
                "Password-protected or encrypted PDF files are not supported.",
            )
        if len(reader.pages) > max_pages:
            raise CandidateDocumentUploadError(
                "too_many_pdf_pages",
                f"The PDF exceeds the {max_pages}-page limit.",
            )
        _validate_pdf_object_graph(reader)

        parts: list[str] = []
        character_count = 0
        for page in reader.pages:
            page_text = page.extract_text() or ""
            character_count += len(page_text)
            if character_count > max_characters:
                raise CandidateDocumentUploadError(
                    "extracted_text_too_large",
                    "The extracted document text exceeds the supported limit.",
                )
            parts.append(page_text)
    except CandidateDocumentUploadError:
        raise
    except Exception as error:
        raise CandidateDocumentUploadError(
            "invalid_pdf",
            "The PDF is malformed or cannot be safely read.",
        ) from error

    return _normalize_extracted_text("\n".join(parts), max_characters)


def _validate_docx_archive(
    raw: bytes,
    *,
    max_entries: int,
    max_expanded_bytes: int,
    max_entry_bytes: int,
) -> None:
    if not zipfile.is_zipfile(BytesIO(raw)):
        raise CandidateDocumentUploadError(
            "invalid_docx_signature",
            "The file content is not a valid DOCX document.",
        )

    try:
        with zipfile.ZipFile(BytesIO(raw)) as archive:
            entries = archive.infolist()
            if len(entries) > max_entries:
                raise CandidateDocumentUploadError(
                    "docx_too_many_entries",
                    "The DOCX package contains too many internal files.",
                )

            names = {entry.filename for entry in entries}
            normalized_names = [
                entry.filename.replace("\\", "/").casefold() for entry in entries
            ]
            if len(normalized_names) != len(set(normalized_names)):
                raise CandidateDocumentUploadError(
                    "duplicate_docx_entry",
                    "The DOCX package contains duplicate internal files.",
                )
            required = {"[Content_Types].xml", "word/document.xml"}
            if not required.issubset(names):
                raise CandidateDocumentUploadError(
                    "invalid_docx_package",
                    "The DOCX package is missing required document content.",
                )

            expanded_size = 0
            for entry in entries:
                normalized_name = entry.filename.replace("\\", "/")
                path = PurePosixPath(normalized_name)
                unix_mode = (entry.external_attr >> 16) & 0xFFFF
                if (
                    normalized_name != entry.filename
                    or path.is_absolute()
                    or ".." in path.parts
                    or (path.parts and path.parts[0].endswith(":"))
                    or stat.S_ISLNK(unix_mode)
                ):
                    raise CandidateDocumentUploadError(
                        "unsafe_docx_path",
                        "The DOCX package contains an unsafe internal path.",
                    )
                if normalized_name.casefold().startswith(
                    ("word/activex/", "word/embeddings/")
                ):
                    raise CandidateDocumentUploadError(
                        "unsafe_docx_content",
                        "Embedded objects and active document content are not "
                        "supported.",
                    )
                if entry.flag_bits & 0x1:
                    raise CandidateDocumentUploadError(
                        "encrypted_docx",
                        "Encrypted DOCX files are not supported.",
                    )
                if entry.file_size > max_entry_bytes:
                    raise CandidateDocumentUploadError(
                        "docx_entry_too_large",
                        "The DOCX package contains an oversized internal file.",
                    )
                expanded_size += entry.file_size
                if expanded_size > max_expanded_bytes:
                    raise CandidateDocumentUploadError(
                        "docx_expansion_too_large",
                        "The expanded DOCX package exceeds the supported limit.",
                    )
                if (
                    entry.file_size > 1_000_000
                    and entry.file_size
                    > max(entry.compress_size, 1) * MAX_DOCX_COMPRESSION_RATIO
                ):
                    raise CandidateDocumentUploadError(
                        "unsafe_docx_compression",
                        "The DOCX package has an unsafe compression ratio.",
                    )

            if any(name.lower().endswith("vbaproject.bin") for name in names):
                raise CandidateDocumentUploadError(
                    "macro_enabled_document",
                    "Macro-enabled Word documents are not supported.",
                )
            if archive.testzip() is not None:
                raise CandidateDocumentUploadError(
                    "corrupt_docx",
                    "The DOCX package is corrupted.",
                )
            for entry in entries:
                lowered_name = entry.filename.lower()
                if not lowered_name.endswith((".xml", ".rels")):
                    continue
                xml_content = archive.read(entry)
                lowered_content = xml_content.lower()
                if b"<!doctype" in lowered_content or b"<!entity" in lowered_content:
                    raise CandidateDocumentUploadError(
                        "unsafe_docx_xml",
                        "The DOCX package contains unsupported XML declarations.",
                    )
                if lowered_name.endswith(".rels"):
                    try:
                        relationships = ElementTree.fromstring(xml_content)
                    except ElementTree.ParseError as error:
                        raise CandidateDocumentUploadError(
                            "invalid_docx_relationships",
                            "The DOCX package contains malformed relationships.",
                        ) from error
                    for relationship in relationships:
                        relationship_type = relationship.attrib.get(
                            "Type", ""
                        ).casefold()
                        target_mode = relationship.attrib.get(
                            "TargetMode", ""
                        ).casefold()
                        if relationship_type.endswith(
                            (
                                "/attachedtemplate",
                                "/oleobject",
                                "/package",
                                "/control",
                                "/afchunk",
                            )
                        ) or (
                            target_mode == "external"
                            and not relationship_type.endswith("/hyperlink")
                        ):
                            raise CandidateDocumentUploadError(
                                "unsafe_docx_relationship",
                                "The DOCX package contains an unsafe external or "
                                "embedded relationship.",
                            )
    except CandidateDocumentUploadError:
        raise
    except Exception as error:
        raise CandidateDocumentUploadError(
            "invalid_docx",
            "The DOCX file is malformed or cannot be safely read.",
        ) from error


def _extract_docx_text(raw: bytes, *, max_characters: int) -> str:
    try:
        document = DocxDocument(BytesIO(raw))
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
    except Exception as error:
        raise CandidateDocumentUploadError(
            "invalid_docx",
            "The DOCX file is malformed or cannot be safely read.",
        ) from error

    return _normalize_extracted_text("\n".join(parts), max_characters)


def extract_cv_text(
    *,
    raw: bytes,
    filename: str,
    declared_content_type: str = "",
    max_pdf_pages: int = DEFAULT_MAX_PDF_PAGES,
    max_docx_entries: int = DEFAULT_MAX_DOCX_ENTRIES,
    max_docx_expanded_bytes: int = DEFAULT_MAX_DOCX_EXPANDED_BYTES,
    max_docx_entry_bytes: int = DEFAULT_MAX_DOCX_ENTRY_BYTES,
    max_extracted_characters: int = DEFAULT_MAX_EXTRACTED_CHARACTERS,
) -> ExtractedCV:
    original_filename = _safe_original_filename(filename)
    extension = PurePosixPath(original_filename).suffix.lower()
    content_type = (
        (declared_content_type or "").split(";", maxsplit=1)[0].strip().lower()
    )

    if extension == ".pdf":
        allowed_types = GENERIC_UPLOAD_CONTENT_TYPES | {PDF_CONTENT_TYPE}
        if content_type not in allowed_types:
            raise CandidateDocumentUploadError(
                "content_type_mismatch",
                "The declared file type does not match a PDF document.",
            )
        text = _extract_pdf_text(
            raw,
            max_pages=max_pdf_pages,
            max_characters=max_extracted_characters,
        )
        return ExtractedCV(
            text=text,
            content_type=PDF_CONTENT_TYPE,
            extension=extension,
        )

    if extension == ".docx":
        allowed_types = GENERIC_UPLOAD_CONTENT_TYPES | {DOCX_CONTENT_TYPE}
        if content_type not in allowed_types:
            raise CandidateDocumentUploadError(
                "content_type_mismatch",
                "The declared file type does not match a DOCX document.",
            )
        _validate_docx_archive(
            raw,
            max_entries=max_docx_entries,
            max_expanded_bytes=max_docx_expanded_bytes,
            max_entry_bytes=max_docx_entry_bytes,
        )
        text = _extract_docx_text(raw, max_characters=max_extracted_characters)
        return ExtractedCV(
            text=text,
            content_type=DOCX_CONTENT_TYPE,
            extension=extension,
        )

    raise CandidateDocumentUploadError(
        "unsupported_extension",
        "Only PDF and DOCX CV files are supported.",
    )


def validate_candidate_cv_upload(
    *,
    uploaded_file,
    max_bytes: int = DEFAULT_MAX_DOCUMENT_BYTES,
) -> ValidatedCVUpload:
    """Read and validate one upload without persisting private bytes."""
    original_filename = _safe_original_filename(getattr(uploaded_file, "name", ""))
    raw = _read_bounded(uploaded_file, max_bytes)
    extracted = extract_cv_text(
        raw=raw,
        filename=original_filename,
        declared_content_type=getattr(uploaded_file, "content_type", ""),
    )
    return ValidatedCVUpload(
        raw=raw,
        original_filename=original_filename,
        extracted=extracted,
        sha256=hashlib.sha256(raw).hexdigest(),
    )


def upload_candidate_cv(
    *,
    candidate: Candidate,
    user: User,
    uploaded_file,
    retention_until: date | None = None,
    max_bytes: int = DEFAULT_MAX_DOCUMENT_BYTES,
) -> CandidateDocument:
    require_organization_object_access(user, candidate)
    validated = validate_candidate_cv_upload(
        uploaded_file=uploaded_file,
        max_bytes=max_bytes,
    )

    document = CandidateDocument(
        candidate=candidate,
        document_type=CandidateDocument.DocumentType.CV,
        original_filename=validated.original_filename,
        file=ContentFile(validated.raw, name=validated.original_filename),
        content_type=validated.extracted.content_type,
        size_bytes=len(validated.raw),
        sha256=validated.sha256,
        extraction_status=CandidateDocument.ExtractionStatus.SUCCEEDED,
        extracted_text=validated.extracted.text,
        extracted_at=timezone.now(),
        retention_until=retention_until,
        uploaded_by=user,
    )

    try:
        with transaction.atomic():
            from organizations.models import Organization

            Organization.objects.select_for_update().get(pk=candidate.organization_id)
            locked_candidate = Candidate.objects.select_for_update().get(
                pk=candidate.pk
            )
            if locked_candidate.status in {
                Candidate.Status.DELETION_REQUESTED,
                Candidate.Status.DELETED,
            }:
                raise CandidateDocumentUploadError(
                    "candidate_unavailable",
                    "A CV cannot be uploaded while candidate deletion is pending.",
                )
            duplicate = (
                CandidateDocument.objects.for_organization(
                    locked_candidate.organization
                )
                .filter(sha256=validated.sha256, deleted_at__isnull=True)
                .select_related("candidate")
                .first()
            )
            if duplicate:
                raise CandidateDocumentDuplicateError(
                    "duplicate_document",
                    "This exact document is already stored for "
                    f"{duplicate.candidate.full_name}.",
                )
            document.candidate = locked_candidate
            document.full_clean()
            document.save()
    except Exception:
        stored_name = document.file.name
        if stored_name and stored_name != validated.original_filename:
            document.file.storage.delete(stored_name)
        raise

    return document


def load_private_candidate_document(
    *,
    document: CandidateDocument,
    user: User,
    max_bytes: int = DEFAULT_MAX_DOCUMENT_BYTES,
) -> PrivateCandidateDocument:
    """Authorize and load one validated private document without exposing its key."""
    require_organization_object_access(user, document)
    if (
        document.deleted_at is not None
        or document.candidate.status == Candidate.Status.DELETED
    ):
        raise CandidateDocumentDeliveryError(
            "document_unavailable",
            "This private document is no longer available.",
        )
    if document.extraction_status != CandidateDocument.ExtractionStatus.SUCCEEDED:
        raise CandidateDocumentDeliveryError(
            "document_unavailable",
            "This private document is not available for download.",
        )
    if (
        document.document_type != CandidateDocument.DocumentType.CV
        or document.content_type not in {PDF_CONTENT_TYPE, DOCX_CONTENT_TYPE}
    ):
        raise CandidateDocumentDeliveryError(
            "invalid_document_metadata",
            "This private document cannot be delivered safely.",
        )

    try:
        filename = _safe_original_filename(document.original_filename)
    except CandidateDocumentUploadError as error:
        raise CandidateDocumentDeliveryError(
            "invalid_document_metadata",
            "This private document cannot be delivered safely.",
        ) from error
    try:
        with document.file.storage.open(document.file.name, "rb") as stored_file:
            content = _read_bounded(stored_file, max_bytes)
    except CandidateDocumentUploadError as error:
        raise CandidateDocumentDeliveryError(
            "document_integrity_failed",
            "This private document failed its integrity check and was not delivered.",
        ) from error
    except (OSError, ValueError) as error:
        raise CandidateDocumentDeliveryError(
            "document_unavailable",
            "This private document is temporarily unavailable.",
        ) from error

    if (
        len(content) > max_bytes
        or document.size_bytes != len(content)
        or not document.sha256
        or hashlib.sha256(content).hexdigest() != document.sha256
    ):
        raise CandidateDocumentDeliveryError(
            "document_integrity_failed",
            "This private document failed its integrity check and was not delivered.",
        )

    return PrivateCandidateDocument(
        content=content,
        filename=filename,
        content_type=document.content_type,
    )

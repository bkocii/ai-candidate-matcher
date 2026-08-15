import zipfile
from datetime import date
from io import BytesIO

import pytest
from django.core.exceptions import PermissionDenied
from django.core.files.uploadedfile import SimpleUploadedFile
from django.db import IntegrityError, transaction
from django.urls import reverse
from django.utils import timezone
from docx import Document as DocxDocument
from pypdf import PdfWriter

from accounts.models import OrganizationMembership, User
from audit.models import AuditEvent
from candidates.documents import (
    DOCX_CONTENT_TYPE,
    CandidateDocumentDuplicateError,
    CandidateDocumentUploadError,
    extract_cv_text,
    load_private_candidate_document,
    upload_candidate_cv,
)
from candidates.models import Candidate, CandidateDocument, CandidateSource
from candidates.services import delete_candidate
from organizations.models import Organization

pytestmark = pytest.mark.django_db


def add_member(
    user: User,
    organization: Organization,
    *,
    role: str = OrganizationMembership.Role.RECRUITER,
) -> None:
    OrganizationMembership.objects.create(
        user=user,
        organization=organization,
        role=role,
    )


def text_pdf(text: str) -> bytes:
    escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream = f"BT /F1 12 Tf 72 720 Td ({escaped}) Tj ET".encode()
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>"
        ),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length "
        + str(len(stream)).encode()
        + b" >>\nstream\n"
        + stream
        + b"\nendstream",
    ]
    pdf = bytearray(b"%PDF-1.4\n")
    offsets = []
    for number, body in enumerate(objects, start=1):
        offsets.append(len(pdf))
        pdf.extend(f"{number} 0 obj\n".encode() + body + b"\nendobj\n")
    xref_offset = len(pdf)
    pdf.extend(f"xref\n0 {len(objects) + 1}\n".encode())
    pdf.extend(b"0000000000 65535 f \n")
    for offset in offsets:
        pdf.extend(f"{offset:010d} 00000 n \n".encode())
    pdf.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode()
    )
    return bytes(pdf)


def docx_bytes(text: str = "Arta Krasniqi — Python and Django") -> bytes:
    document = DocxDocument()
    document.add_heading("Curriculum Vitae", level=1)
    document.add_paragraph(text)
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Experience"
    table.cell(0, 1).text = "Five years"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def pdf_upload(content: bytes | None = None) -> SimpleUploadedFile:
    return SimpleUploadedFile(
        "candidate.pdf",
        content or text_pdf("Candidate Python Django PostgreSQL"),
        content_type="application/pdf",
    )


def pdf_with_active_content(kind: str) -> bytes:
    writer = PdfWriter()
    writer.append(BytesIO(text_pdf("Synthetic candidate text")))
    if kind == "javascript":
        writer.add_js("app.alert('synthetic')")
    else:
        writer.add_attachment("payload.txt", b"synthetic attachment")
    output = BytesIO()
    writer.write(output)
    return output.getvalue()


def docx_with_relationship(*, relationship_type: str, target_mode: str) -> bytes:
    output = BytesIO()
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", "<Types />")
        archive.writestr("word/document.xml", "<document />")
        archive.writestr(
            "word/_rels/document.xml.rels",
            (
                '<Relationships xmlns="http://schemas.openxmlformats.org/'
                'package/2006/relationships">'
                f'<Relationship Id="rId1" Type="{relationship_type}" '
                f'Target="https://example.test/resource" '
                f'TargetMode="{target_mode}" />'
                "</Relationships>"
            ),
        )
    return output.getvalue()


def test_extracts_text_from_pdf_and_docx() -> None:
    pdf = extract_cv_text(
        raw=text_pdf("Arta Krasniqi Python Django"),
        filename="arta.pdf",
        declared_content_type="application/pdf",
    )
    docx = extract_cv_text(
        raw=docx_bytes(),
        filename="arta.docx",
        declared_content_type=DOCX_CONTENT_TYPE,
    )

    assert pdf.content_type == "application/pdf"
    assert "Python Django" in pdf.text
    assert docx.content_type == DOCX_CONTENT_TYPE
    assert "Arta Krasniqi" in docx.text
    assert "Five years" in docx.text


@pytest.mark.parametrize(
    ("filename", "content_type", "raw", "code"),
    [
        ("candidate.doc", "application/msword", b"legacy", "unsupported_extension"),
        ("candidate.pdf", "text/plain", b"%PDF-fake", "content_type_mismatch"),
        ("candidate.pdf", "application/pdf", b"not a pdf", "invalid_pdf_signature"),
        (
            "candidate.docx",
            DOCX_CONTENT_TYPE,
            b"not a zip package",
            "invalid_docx_signature",
        ),
    ],
)
def test_rejects_unsupported_mismatched_or_disguised_files(
    filename: str,
    content_type: str,
    raw: bytes,
    code: str,
) -> None:
    with pytest.raises(CandidateDocumentUploadError) as captured:
        extract_cv_text(
            raw=raw,
            filename=filename,
            declared_content_type=content_type,
        )

    assert captured.value.code == code


def test_rejects_encrypted_or_textless_pdf() -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    writer.encrypt("synthetic-password")
    encrypted = BytesIO()
    writer.write(encrypted)

    with pytest.raises(CandidateDocumentUploadError) as captured:
        extract_cv_text(raw=encrypted.getvalue(), filename="private.pdf")
    assert captured.value.code == "encrypted_pdf"

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    blank = BytesIO()
    writer.write(blank)
    with pytest.raises(CandidateDocumentUploadError) as captured:
        extract_cv_text(raw=blank.getvalue(), filename="scan.pdf")
    assert captured.value.code == "no_extractable_text"


@pytest.mark.parametrize("kind", ["javascript", "attachment"])
def test_pdf_active_content_and_embedded_files_are_rejected(kind: str) -> None:
    with pytest.raises(CandidateDocumentUploadError) as captured:
        extract_cv_text(
            raw=pdf_with_active_content(kind),
            filename="candidate.pdf",
            declared_content_type="application/pdf",
        )

    assert captured.value.code == "unsafe_pdf_content"


def test_pdf_page_and_extracted_text_limits_are_enforced() -> None:
    raw = text_pdf("Candidate text longer than the configured limit")

    with pytest.raises(CandidateDocumentUploadError) as captured:
        extract_cv_text(raw=raw, filename="candidate.pdf", max_pdf_pages=0)
    assert captured.value.code == "too_many_pdf_pages"

    with pytest.raises(CandidateDocumentUploadError) as captured:
        extract_cv_text(raw=raw, filename="candidate.pdf", max_extracted_characters=5)
    assert captured.value.code == "extracted_text_too_large"


def test_docx_archive_expansion_limits_are_enforced() -> None:
    output = BytesIO()
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", "<Types />")
        archive.writestr("word/document.xml", " " * 2_000_000)

    with pytest.raises(CandidateDocumentUploadError) as captured:
        extract_cv_text(
            raw=output.getvalue(),
            filename="candidate.docx",
            max_docx_expanded_bytes=1_000_000,
        )

    assert captured.value.code in {
        "docx_entry_too_large",
        "docx_expansion_too_large",
        "unsafe_docx_compression",
    }


def test_docx_with_entity_declarations_is_rejected() -> None:
    output = BytesIO()
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", "<Types />")
        archive.writestr(
            "word/document.xml",
            '<!DOCTYPE document [<!ENTITY cv "private">]><document>&cv;</document>',
        )

    with pytest.raises(CandidateDocumentUploadError) as captured:
        extract_cv_text(raw=output.getvalue(), filename="candidate.docx")

    assert captured.value.code == "unsafe_docx_xml"


def test_docx_unsafe_external_relationship_is_rejected() -> None:
    relationship_type = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/"
        "attachedTemplate"
    )

    with pytest.raises(CandidateDocumentUploadError) as captured:
        extract_cv_text(
            raw=docx_with_relationship(
                relationship_type=relationship_type,
                target_mode="External",
            ),
            filename="candidate.docx",
            declared_content_type=DOCX_CONTENT_TYPE,
        )

    assert captured.value.code == "unsafe_docx_relationship"


def test_docx_symlink_entry_is_rejected() -> None:
    output = BytesIO()
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", "<Types />")
        archive.writestr("word/document.xml", "<document />")
        symlink = zipfile.ZipInfo("word/media/link.png")
        symlink.create_system = 3
        symlink.external_attr = (0o120777 << 16) | 0xA1ED0000
        archive.writestr(symlink, "../../../private")

    with pytest.raises(CandidateDocumentUploadError) as captured:
        extract_cv_text(raw=output.getvalue(), filename="candidate.docx")

    assert captured.value.code == "unsafe_docx_path"


def test_filename_with_bidirectional_control_is_rejected() -> None:
    with pytest.raises(CandidateDocumentUploadError) as captured:
        extract_cv_text(
            raw=text_pdf("Synthetic candidate"),
            filename="candidate\N{RIGHT-TO-LEFT OVERRIDE}.pdf",
            declared_content_type="application/pdf",
        )

    assert captured.value.code == "invalid_filename"


def test_upload_stores_private_metadata_hash_and_extracted_text(
    settings,
    tmp_path,
) -> None:
    settings.MEDIA_ROOT = tmp_path
    user = User.objects.create_user(username="recruiter")
    organization = Organization.objects.create(name="Northstar", slug="northstar")
    add_member(user, organization)
    candidate = Candidate.objects.create(
        organization=organization,
        full_name="Arta Krasniqi",
    )

    document = upload_candidate_cv(
        candidate=candidate,
        user=user,
        uploaded_file=SimpleUploadedFile(
            "folder/Arta CV.docx",
            docx_bytes(),
            content_type=DOCX_CONTENT_TYPE,
        ),
        retention_until=date(2027, 8, 10),
    )

    assert document.original_filename == "Arta CV.docx"
    assert document.document_type == CandidateDocument.DocumentType.CV
    assert document.extraction_status == CandidateDocument.ExtractionStatus.SUCCEEDED
    assert "Python and Django" in document.extracted_text
    assert document.extracted_at is not None
    assert document.extraction_error_code == ""
    assert document.size_bytes == len(docx_bytes())
    assert len(document.sha256) == 64
    assert document.uploaded_by == user
    assert document.retention_until == date(2027, 8, 10)
    assert document.file.name.startswith(
        f"candidate_documents/{organization.pk}/{candidate.pk}/"
    )
    assert "Arta" not in document.file.name
    assert document.file.storage.exists(document.file.name)


def test_duplicate_hash_check_is_organization_scoped(settings, tmp_path) -> None:
    settings.MEDIA_ROOT = tmp_path
    user = User.objects.create_user(username="recruiter")
    first = Organization.objects.create(name="First", slug="first")
    second = Organization.objects.create(name="Second", slug="second")
    add_member(user, first)
    add_member(user, second)
    first_candidate = Candidate.objects.create(organization=first, full_name="First")
    second_candidate = Candidate.objects.create(
        organization=second,
        full_name="Second",
    )
    raw = text_pdf("Same synthetic CV")

    upload_candidate_cv(
        candidate=first_candidate,
        user=user,
        uploaded_file=pdf_upload(raw),
    )
    with pytest.raises(CandidateDocumentDuplicateError) as captured:
        upload_candidate_cv(
            candidate=first_candidate,
            user=user,
            uploaded_file=pdf_upload(raw),
        )
    assert captured.value.code == "duplicate_document"

    second_document = upload_candidate_cv(
        candidate=second_candidate,
        user=user,
        uploaded_file=pdf_upload(raw),
    )
    assert second_document.candidate == second_candidate
    assert CandidateDocument.objects.count() == 2


def test_upload_service_requires_membership_before_storing(settings, tmp_path) -> None:
    settings.MEDIA_ROOT = tmp_path
    outsider = User.objects.create_user(username="outsider")
    organization = Organization.objects.create(name="Northstar", slug="northstar")
    candidate = Candidate.objects.create(
        organization=organization,
        full_name="Candidate",
    )

    with pytest.raises(PermissionDenied):
        upload_candidate_cv(
            candidate=candidate,
            user=outsider,
            uploaded_file=pdf_upload(),
        )

    assert not CandidateDocument.objects.exists()
    assert not tmp_path.exists() or not any(tmp_path.rglob("*"))


def test_upload_service_rejects_deleted_candidate(settings, tmp_path) -> None:
    settings.MEDIA_ROOT = tmp_path
    user = User.objects.create_user(username="recruiter")
    organization = Organization.objects.create(name="Northstar", slug="northstar")
    add_member(user, organization)
    candidate = Candidate.objects.create(
        organization=organization,
        full_name="Deleted candidate #1",
        status=Candidate.Status.DELETED,
        deletion_requested_at=timezone.now(),
        deleted_at=timezone.now(),
    )

    with pytest.raises(CandidateDocumentUploadError) as captured:
        upload_candidate_cv(
            candidate=candidate,
            user=user,
            uploaded_file=pdf_upload(),
        )

    assert captured.value.code == "candidate_unavailable"
    assert not CandidateDocument.objects.exists()
    assert not tmp_path.exists() or not any(tmp_path.rglob("*"))


def test_failed_extraction_does_not_store_file(settings, tmp_path) -> None:
    settings.MEDIA_ROOT = tmp_path
    user = User.objects.create_user(username="recruiter")
    organization = Organization.objects.create(name="Northstar", slug="northstar")
    add_member(user, organization)
    candidate = Candidate.objects.create(
        organization=organization,
        full_name="Candidate",
    )

    with pytest.raises(CandidateDocumentUploadError):
        upload_candidate_cv(
            candidate=candidate,
            user=user,
            uploaded_file=pdf_upload(b"not a pdf"),
        )

    assert not CandidateDocument.objects.exists()
    assert not tmp_path.exists() or not any(tmp_path.rglob("*"))


def test_successful_extraction_requires_text_and_timestamp() -> None:
    organization = Organization.objects.create(name="Northstar", slug="northstar")
    candidate = Candidate.objects.create(
        organization=organization,
        full_name="Candidate",
    )

    with pytest.raises(IntegrityError):
        with transaction.atomic():
            CandidateDocument.objects.create(
                candidate=candidate,
                original_filename="candidate.pdf",
                file="candidate_documents/candidate.pdf",
                extraction_status=CandidateDocument.ExtractionStatus.SUCCEEDED,
            )


def test_recruiter_can_upload_cv_and_view_only_safe_metadata(
    client,
    settings,
    tmp_path,
) -> None:
    settings.MEDIA_ROOT = tmp_path
    user = User.objects.create_user(username="recruiter")
    organization = Organization.objects.create(name="Northstar", slug="northstar")
    add_member(user, organization)
    candidate = Candidate.objects.create(
        organization=organization,
        full_name="Arta Krasniqi",
    )
    client.force_login(user)

    response = client.post(
        reverse(
            "candidates:candidate-cv-upload",
            kwargs={
                "organization_slug": organization.slug,
                "candidate_id": candidate.pk,
            },
        ),
        {
            "cv_file": SimpleUploadedFile(
                "arta.docx",
                docx_bytes("PRIVATE-CV-TEXT Python Django"),
                content_type=DOCX_CONTENT_TYPE,
            ),
            "retention_until": "2027-08-10",
        },
    )

    document = CandidateDocument.objects.get()
    assert response.status_code == 302
    assert document.candidate == candidate
    assert "PRIVATE-CV-TEXT" in document.extracted_text

    detail = client.get(response.url)
    content = detail.content.decode()
    assert detail.status_code == 200
    assert "arta.docx" in content
    assert "Succeeded" in content
    assert "PRIVATE-CV-TEXT" not in content
    assert document.file.name not in content
    assert "Download original" in content


def test_recruiter_can_privately_download_original_cv(
    client,
    settings,
    tmp_path,
) -> None:
    settings.MEDIA_ROOT = tmp_path
    user = User.objects.create_user(username="recruiter")
    organization = Organization.objects.create(name="Northstar", slug="northstar")
    add_member(user, organization)
    candidate = Candidate.objects.create(
        organization=organization,
        full_name="Arta Krasniqi",
    )
    raw = text_pdf("PRIVATE SYNTHETIC CV")
    document = upload_candidate_cv(
        candidate=candidate,
        user=user,
        uploaded_file=SimpleUploadedFile(
            "Arta CV.pdf",
            raw,
            content_type="application/pdf",
        ),
    )
    client.force_login(user)

    response = client.get(
        reverse(
            "candidates:candidate-document-download",
            args=[organization.slug, candidate.pk, document.pk],
        )
    )

    assert response.status_code == 200
    assert b"".join(response.streaming_content) == raw
    assert response["Content-Type"] == "application/pdf"
    assert "attachment" in response["Content-Disposition"]
    assert "Arta CV.pdf" in response["Content-Disposition"]
    assert document.file.name not in response["Content-Disposition"]
    assert response["Cache-Control"] == "private, no-store, max-age=0"
    assert response["Pragma"] == "no-cache"
    assert response["X-Content-Type-Options"] == "nosniff"
    assert response["Content-Security-Policy"] == "sandbox"
    assert response["Cross-Origin-Resource-Policy"] == "same-origin"
    assert response["Referrer-Policy"] == "no-referrer"
    assert response["X-Robots-Tag"] == "noindex, nofollow, noarchive"
    event = AuditEvent.objects.get(
        action=AuditEvent.Action.CANDIDATE_DOCUMENT_DOWNLOADED
    )
    assert event.organization == organization
    assert event.actor == user
    assert event.object_type == AuditEvent.ObjectType.CANDIDATE_DOCUMENT
    assert event.object_id == document.pk


def test_private_download_repeats_service_authorization(settings, tmp_path) -> None:
    settings.MEDIA_ROOT = tmp_path
    owner = User.objects.create_user(username="owner")
    outsider = User.objects.create_user(username="outsider")
    organization = Organization.objects.create(name="Northstar", slug="northstar")
    add_member(owner, organization)
    candidate = Candidate.objects.create(
        organization=organization,
        full_name="Candidate",
    )
    document = upload_candidate_cv(
        candidate=candidate,
        user=owner,
        uploaded_file=pdf_upload(),
    )

    with pytest.raises(PermissionDenied):
        load_private_candidate_document(document=document, user=outsider)


def test_private_download_hides_cross_tenant_and_candidate_mismatch(
    client,
    settings,
    tmp_path,
) -> None:
    settings.MEDIA_ROOT = tmp_path
    user = User.objects.create_user(username="recruiter")
    hidden_owner = User.objects.create_user(username="hidden-owner")
    visible = Organization.objects.create(name="Visible", slug="visible")
    hidden = Organization.objects.create(name="Hidden", slug="hidden")
    add_member(user, visible)
    add_member(hidden_owner, hidden)
    visible_candidate = Candidate.objects.create(
        organization=visible,
        full_name="Visible Candidate",
    )
    hidden_candidate = Candidate.objects.create(
        organization=hidden,
        full_name="Hidden Candidate",
    )
    hidden_document = upload_candidate_cv(
        candidate=hidden_candidate,
        user=hidden_owner,
        uploaded_file=pdf_upload(),
    )
    client.force_login(user)

    cross_tenant = client.get(
        reverse(
            "candidates:candidate-document-download",
            args=[hidden.slug, hidden_candidate.pk, hidden_document.pk],
        )
    )
    wrong_candidate = client.get(
        reverse(
            "candidates:candidate-document-download",
            args=[visible.slug, visible_candidate.pk, hidden_document.pk],
        )
    )

    assert cross_tenant.status_code == 404
    assert wrong_candidate.status_code == 404


def test_private_download_refuses_changed_storage_bytes(
    client,
    settings,
    tmp_path,
) -> None:
    settings.MEDIA_ROOT = tmp_path
    user = User.objects.create_user(username="recruiter")
    organization = Organization.objects.create(name="Northstar", slug="northstar")
    add_member(user, organization)
    candidate = Candidate.objects.create(
        organization=organization,
        full_name="Candidate",
    )
    document = upload_candidate_cv(
        candidate=candidate,
        user=user,
        uploaded_file=pdf_upload(),
    )
    with document.file.storage.open(document.file.name, "wb") as stored_file:
        stored_file.write(b"changed private bytes")
    client.force_login(user)

    response = client.get(
        reverse(
            "candidates:candidate-document-download",
            args=[organization.slug, candidate.pk, document.pk],
        ),
        follow=True,
    )

    assert response.status_code == 200
    assert response.redirect_chain
    assert "failed its integrity check" in response.content.decode()
    assert not getattr(response, "streaming", False)


@pytest.mark.parametrize("route_name", ["candidate-detail", "candidate-cv-upload"])
def test_document_routes_hide_cross_organization_candidates(client, route_name) -> None:
    user = User.objects.create_user(username="recruiter")
    visible = Organization.objects.create(name="Visible", slug="visible")
    hidden = Organization.objects.create(name="Hidden", slug="hidden")
    add_member(user, visible)
    hidden_candidate = Candidate.objects.create(
        organization=hidden,
        full_name="Hidden Candidate",
    )
    client.force_login(user)

    response = client.get(
        reverse(
            f"candidates:{route_name}",
            kwargs={
                "organization_slug": hidden.slug,
                "candidate_id": hidden_candidate.pk,
            },
        )
    )

    assert response.status_code == 404
    assert "Hidden Candidate" not in response.content.decode()


def test_upload_view_returns_safe_parser_error(client) -> None:
    user = User.objects.create_user(username="recruiter")
    organization = Organization.objects.create(name="Northstar", slug="northstar")
    add_member(user, organization)
    candidate = Candidate.objects.create(
        organization=organization,
        full_name="Candidate",
    )
    client.force_login(user)

    response = client.post(
        reverse(
            "candidates:candidate-cv-upload",
            kwargs={
                "organization_slug": organization.slug,
                "candidate_id": candidate.pk,
            },
        ),
        {
            "cv_file": SimpleUploadedFile(
                "broken.pdf",
                b"%PDF-malformed parser internals",
                content_type="application/pdf",
            )
        },
    )

    content = response.content.decode()
    assert response.status_code == 200
    assert "malformed or cannot be safely read" in content
    assert "parser internals" not in content
    assert not CandidateDocument.objects.exists()


def test_django_admin_cannot_bypass_validated_document_creation(client) -> None:
    superuser = User.objects.create_superuser(
        username="operator",
        password="synthetic-password",
    )
    client.force_login(superuser)

    response = client.get(reverse("admin:candidates_candidatedocument_add"))

    assert response.status_code == 403


def test_candidate_deletion_requires_request_then_admin_purge(
    client, settings, tmp_path
) -> None:
    settings.MEDIA_ROOT = tmp_path
    user = User.objects.create_user(username="recruiter")
    organization = Organization.objects.create(name="Northstar", slug="northstar")
    add_member(user, organization, role=OrganizationMembership.Role.ADMIN)
    candidate = Candidate.objects.create(
        organization=organization,
        full_name="Arta Krasniqi",
        email="arta@example.test",
        phone="+383 44 111 222",
        location="Prishtina",
        created_by=user,
    )
    CandidateSource.objects.create(
        candidate=candidate,
        source_type=CandidateSource.SourceType.MANUAL_ENTRY,
        source_name="Synthetic source",
        source_reference="PRIVATE-REF",
        permission_notes="Private candidate notes",
        recorded_by=user,
    )
    document = upload_candidate_cv(
        candidate=candidate,
        user=user,
        uploaded_file=pdf_upload(),
    )
    stored_name = document.file.name
    storage = document.file.storage
    client.force_login(user)
    route = reverse(
        "candidates:candidate-delete",
        args=[organization.slug, candidate.pk],
    )

    request_confirmation = client.get(route)
    request_response = client.post(route)

    candidate.refresh_from_db()
    assert request_confirmation.status_code == 200
    assert "does not purge data" in request_confirmation.content.decode()
    assert request_response.status_code == 302
    assert candidate.status == Candidate.Status.DELETION_REQUESTED
    assert candidate.status_before_deletion_request == Candidate.Status.ACTIVE
    assert candidate.deletion_requested_by == user
    assert CandidateSource.objects.filter(candidate=candidate).exists()
    assert CandidateDocument.objects.filter(candidate=candidate).exists()
    assert storage.exists(stored_name)

    purge_route = reverse(
        "candidates:candidate-delete-execute",
        args=[organization.slug, candidate.pk],
    )
    purge_confirmation = client.get(purge_route)
    response = client.post(purge_route)

    candidate.refresh_from_db()
    assert purge_confirmation.status_code == 200
    assert "permanently purge" in purge_confirmation.content.decode().lower()
    assert response.status_code == 302
    assert candidate.status == Candidate.Status.DELETED
    assert candidate.full_name == f"Deleted candidate #{candidate.pk}"
    assert candidate.email == ""
    assert candidate.phone == ""
    assert candidate.location == ""
    assert candidate.deletion_requested_at is not None
    assert candidate.deleted_at is not None
    assert not CandidateSource.objects.filter(candidate=candidate).exists()
    assert not CandidateDocument.objects.filter(candidate=candidate).exists()
    assert not storage.exists(stored_name)

    listing = client.get(reverse("candidates:candidate-list", args=[organization.slug]))
    detail = client.get(
        reverse("candidates:candidate-detail", args=[organization.slug, candidate.pk])
    )
    assert list(listing.context["page"].object_list) == []
    assert detail.status_code == 404
    dashboard = client.get(
        reverse("organizations:organization-dashboard", args=[organization.slug])
    )
    assert dashboard.context["active_candidate_count"] == 0
    assert list(
        AuditEvent.objects.filter(object_id=candidate.pk).values_list(
            "action", flat=True
        )
    ) == [
        AuditEvent.Action.CANDIDATE_DELETED,
        AuditEvent.Action.CANDIDATE_DELETION_REQUESTED,
    ]


def test_candidate_delete_service_repeats_permission_check() -> None:
    owner = User.objects.create_user(username="owner")
    outsider = User.objects.create_user(username="outsider")
    organization = Organization.objects.create(name="Northstar", slug="northstar")
    add_member(owner, organization)
    candidate = Candidate.objects.create(
        organization=organization,
        full_name="Protected Candidate",
    )

    with pytest.raises(PermissionDenied):
        delete_candidate(candidate=candidate, user=outsider)

    candidate.refresh_from_db()
    assert candidate.status == Candidate.Status.ACTIVE


def test_candidate_delete_route_hides_cross_organization_candidate(client) -> None:
    user = User.objects.create_user(username="recruiter")
    visible = Organization.objects.create(name="Visible", slug="visible")
    hidden = Organization.objects.create(name="Hidden", slug="hidden")
    add_member(user, visible)
    candidate = Candidate.objects.create(
        organization=hidden,
        full_name="Hidden Candidate",
    )
    client.force_login(user)

    response = client.post(
        reverse(
            "candidates:candidate-delete",
            args=[visible.slug, candidate.pk],
        )
    )

    assert response.status_code == 404
    candidate.refresh_from_db()
    assert candidate.status == Candidate.Status.ACTIVE
